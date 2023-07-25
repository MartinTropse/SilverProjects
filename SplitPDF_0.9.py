# -*- coding: utf-8 -*-
"""
Created on Tue Dec  3 13:51:52 2019
@author: MartinTropse
"""

import pytesseract 
from wand.image import Image as wi
from PIL import Image
import io
import re
import docx
import time
import os
import subprocess
import pickle

pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
os.chdir(r"C:\myGit\TestBox")
myPdf="C:/myGit/TestBox/Objektskatalog.pdf"

machineList= []
extracted_text = []    
imgBlobs = []
pdf = wi(filename=myPdf, resolution = 300)
pdfImg = pdf.convert('jpeg')
pdf.close()

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

for img in pdfImg.sequence:
    page=wi(image=img) #This opens an actual image
    imgBlobs.append(page.make_blob('jpeg')) #This object contains byte data, make_blob cause the conversion    
    #The bytes object is the converted back to image before being passed to pytesseract    

for imgBlob in imgBlobs:
    im = Image.open(io.BytesIO(imgBlob)) #"im" is an PIL, JpegImageFile which is assumes is the prefered object of tessearct
    text = pytesseract.image_to_string(im, lang='swe')
    extracted_text.append(text)    



"""
Create and load the pickleObject from the text extraction
"""
pickle_out = open("PdfText.pickle", "wb")
pickle.dump(extracted_text, pickle_out)
pickle_out.close()
pickle_in = open("PdfText.pickle", "rb")
extracted_text = pickle.load(pickle_in)

#firstText=extracted_text[0:496]

pos = -1
pdfString=str(extracted_text)
pat = re.compile(r'Nr Namn Klass\\n\d{6}', re.M) # pattern: Upercase, then anything that is not in (.!?), then one of them

pdfaList=pat.split(pdfString)
identifer = re.findall(pat, pdfString)
headTitle = re.compile(r"(\w+\s)+\d")

ntrBed=re.compile(r"(?<=\\nNaturvärdesbedömning\\n).+(?=(\\n){1,3}Naturvärdesbeskrivning(\\n){1,3})")
ntrBed2=re.compile(r"(?<=- Naturvärdesbedömning).+(?=(\\n){1,3}Naturvärdesbeskrivning(\\n){1,3})")
ntrBed3=re.compile("(?<=Naturvärdesbedömning).+(?= Naturvärdesbeskrivning|Naturvärdesbeskrivning)")

ntrBsk=re.compile(r"(?<=\\nNaturvärdesbeskrivning\\n).+?(?=(\\n){1,3}Åtgärdsbehov(\\n){1,3}| 'Åtgärdsbehov|Åtgärdsbehov|Atgärdsbehov)")

ntrBsk2=re.compile(r"(?<=Naturvärdesbeskrivning\\n\\n).+?(?=(\\n){1,3}Åtgärdsbehov(\\n){1,3}| 'Åtgärdsbehov|Åtgärdsbehov)") 
åtBhv=re.compile(r"(?<=Åtgärdsbehov\\n|Atgärdsbehov\\n).+(?=', '\\n\\n|\\n\\n|\.', ')")
åtBhv3=re.compile(r"(?<=Åtgärdsbehov',).+(?=', '\\n\\n|\\n\\n|', ')")
åtBhv2=re.compile(r"(?<=Åtgärdsbehov\\n|Atgärdsbehov\\n).??")


checkStr=docStr

hit1=re.search(ntrBed2, checkStr)
hit2=re.search(ntrBsk2, checkStr)
hit3=re.search(ntrBed3, checkStr)
hit4=re.search(ntrBsk, checkStr)
hit5=re.search(åtBhv, checkStr)
hit6=re.search(åtBhv2, checkStr)
hit6=re.search(åtBhv3, checkStr)

hit1.group(0)
hit2.group(0)
hit6.group(0)
#chkLine="\\nNaturvärdesbeskrivning\\n\\nOrientering\\n\\nPassdalsån ligger väster om Ramsdal i sydöstra delen av Söderköping kommun.\\nVattendraget saknar namn på lantmäteriets kartor, men kallas Passdalsån och\\nKorsnäsbäcken. Passdalsån bedöms vara det mest vedertagna namnet och används\\ndärför här. Bäcken rinner norrut från Nöstebosjön till kusten där den mynnar i\\nGropviken. Den är 4,2 kilometer lång och har en fallhöjd på 17,6 meter.\\nVattendragsbeskrivning\\n\\nBäcken har huvudsakligen ett ringlande lopp, men i de nedre delarna finns flera\\nmeandrande partier. I de centrala och övre delarna finns dessutom en hel del uträtade\\nsträckor. 58 % av vattendraget är välskuggat och 28 % relativt välskuggat. Partier\\nmed sämre skuggning finns främst i området närmast uppströms gården Korsnäs och\\ninom de översta delarna kring Nöstebosjöns utlopp. Död ved finns bara i mindre\\nomfattning i ån. Medelbredden är oftast mellan 2,5 och 4 meter och medeldjupet\\nnågon till några decimeter.\\n\\nVattenmiljön utgörs till största delen av lugnflytande till svagt strömmande sträckor,\\nvilka dominerar inom 54 respektive 39 % av vattendraget. Strömmande partier\\nförekommer främst inom området 0,8 till 2,3 kilometer uppströms utloppet i\\nGropviken och dominerar inom 7 % av bäcken. Finsediment är det vanligaste\\nbottenmaterialet och dominerar inom 68 % av vattendraget. Grövre bottenmaterial i\\nform av sand, grus, sten och block finns främst inom de svagt strömmande och\\nströmmande partierna. Sten är vanligast och dominerar inom 26 % av bäcken, men\\näven sandbottnar är ställvis vanliga.\\n\\nDet finns några särskilt värdefulla avsnitt i bäcken. Det mest intressanta området är\\n720 meter långt och ligger 0,8 till 1,5 kilometer uppströms utloppet i Gropviken. Här\\nrinner bäcken i en ravin och har bitvis ett meandrande lopp. Inga betydande fysiska\\ningrepp har gjorts inom området. I de nedre delarna är ravinen grund och bred och\\n\\n', 'bäcken meandrar över ett brett flodplan. Längre upp blir flodplanet småkulligt och allt\\nsmalare och ravinen successivt allt djupare. Här finns också flera branta strandbrinkar\\nmed blottlagt finkornigt material.\\n\\nRavinen är bevuxen med ask med inslag av bland annat klibbal, lönn, hassel och ek (se\\näven annat NVP-objekt). I de övre delarna finns dessutom en hel del gran och flera\\ngrova granlågor. Delar av ravinen betas. Hela området är en skoglig nyckelbiotop.\\nVattenmiljön är varierande och består omväxlande av strömsträckor, nackar och\\nhöljor. De nedre delarna är något flackare och domineras av svagt strömmande vatten,\\nmen med ett stort inslag av strömmande partier. Längre upp tilltar lutningen något\\noch här dominerar det strömmade vattnet. Sten och grus är de dominerande\\nbottenmaterialen, men det finns även en hel del finsediment, sand och block. Området\\när mycket lämpligt för öring och utgör kärnområde för bäckens havsvandrande\\nbestånd.\\n\\nUpp- och nedströms nyckelbiotopen ansluter två andra mycket värdefulla avsnitt. Den\\nnedre är 200 meter och den övre 500 meter lång. Det nedre området utgörs av en\\nmeandersträcka med svagt strömmande vatten. Bäcken rinner över en öppen hagmark\\noch meandrar över ett grunt nedskuret flodplan av varierande bredd. Sträckan är\\npåverkad av rensningar, men området har kvar sin naturliga karaktär. Flera gamla\\nmeanderslingor kan ses vid sidan och bäcken och åtminstone en av dessa bedöms ha\\nskurits av vid rensningsarbeten. Botten och stranden är påverkad av trampskador från\\ndjuren som betar hagmarken.\\n\\nDet övre området utgörs av en ravinsträcka där bäcken rinner fram över ett bitvis\\nsmåkulligt flodplan. Den grunda ravinen har ett värdefullt trädskikt och är en skoglig\\nnyckelbiotop. Grov ek och klibbal och en hel del hassel växer längs bäcken och\\nskuggningen är varierande med både välskuggade och något sämre skuggade partier.\\n\\n', 'Bäcken har ett ringlande lopp som bitvis gränsar till meandrande och vattnet är\\nmestadels svagt strömmande. Atminstone de övre delarna är påverkade av rensningar,\\nmen området har kvar sin naturliga karaktär. Längst upp finns resterna av Passdals\\nkvarn.\\n\\nVäxter och djur\\n\\nVid karteringen dominerades vattenvegetationen av övervattensväxter och\\nvattenmossor, men förekomsten var sällan särskilt riklig. Skogssäv och näckmossa\\nvar de vanligaste arterna, men även bladvass, mannagräs, rörflen, sjöfräken, igelknopp\\noch lånke hörde ställvis till de dominerande arterna. Näckmossa fanns främst inom\\npartier med något högre vattenhastighet och grövre bottenmaterial medan bladvass\\nföreträdelsevis påträffades i området närmast utloppet i Gropviken samt inom de övre\\ndelarna. Andra arter som noterades var svalting, möja, starr, äkta förgätmigej, andmat,\\ngul svärdslilja, kabbleka, strandlysing, bäcklav, hästskräppa, veketåg, gul näckros,\\nbredkaveldun, fackelblomster och fintrådiga grönalger.\\n\\nBottenfaunan i bäcken provtogs 1995 och 2000 inom ramen för riksinventeringen av\\nsjöar och vattendrag (Institutionen för miljöanalys 2007). Fiskfaunan är artrik och\\ndomineras helt av öring, men även mört, gädda, abborre, nejonöga, gers, id, ål och\\nlake har påträffats. Arterna har fångats vid de trettiotalet elfisken som utförts i bäcken\\nunder perioden 1991 till 2006 (Fiskeriverket 2007, Gustafsson 2007 och 2008 samt\\nGöthberg & Karlsson 2006). Större delen av provfiskena har utförts på två lokaler\\nbelägna omkring 0,9 respektive 1,9 kilometer från bäckens mynning i Gropviken.\\n\\nUtöver detta har några provfisken utförts mellan dessa lokaler. Dessutom har ett fiske\\nutförts vid Stämsel cirka 0,5 kilometer nedströms bäckens utlopp ur Nöstebosjön.\\nÖringen är havsvandrande och har fångats på alla lokaler utom den översta.\\n\\n', 'Åtgärdsbehov\\n\\n"

iPos = 0
pfPos = 1
os.chdir(r"C:\myGit\TestBox")


for _ in pdfaList[1:]:
    myDoc = docx.Document("5GreatDocument.docx")
    docStr=pdfaList[pfPos]
    docStr=re.sub("den\s\d{1,2}\s[a-z]{1,14}\s\d{4}\sSida\s\d{1,3}\sav\s\d{1,3}", "", docStr)
    docStr=re.sub('—','-', docStr)
    docStr=re.sub(r'\bla\\n\\n', "1a\\n\\n", docStr)
    Title_0=re.search(headTitle, docStr)
    klass=Title_0.group(0)[-1]
    Title=Title_0.group(0)[:-1]
    myDoc.paragraphs[1].runs[0].text = Title_0.group(0)[:-1] #Correct
    bdm=re.search(ntrBed, docStr)
    if not bdm:
        bdm = re.search(ntrBed2, docStr)
    if not bdm:
        bdm = re.search(ntrBed3, docStr) #A succesion of regex searches if the prior one doesnt hits
    myDoc.paragraphs[5].runs[2].text = bdm.group(0).replace(r"\n", " ")+"\n"
    if myDoc.paragraphs[5].runs[2].text[0] == " ":
        myDoc.paragraphs[5].runs[2].text = myDoc.paragraphs[5].runs[2].text[1:]    
    bsk=re.search(ntrBsk, docStr)
    if not bsk:
        bsk = re.search(ntrBsk2, docStr) 
    myDoc.paragraphs[5].runs[6].text = bsk.group(0).replace(r"\n", " ")+"\n"
    if myDoc.paragraphs[5].runs[6].text[0] == " ":
       myDoc.paragraphs[5].runs[6].text = myDoc.paragraphs[5].runs[6].text[1:]    
    åtb=re.search(åtBhv, docStr)
    if not åtb:
        åtb = re.search(åtBhv3, docStr)
    if not åtb:
        åtb = re.search(åtBhv2, docStr) #This some insanity code to deal with sections that lacks a text
        if åtb:
               åtb = re.search("   ", "   ") #So if the second regex hit. Replace åtb.group(0) with a "fake string"
    myDoc.paragraphs[5].runs[10].text = åtb.group(0).replace(r"\n", " ") 
    if myDoc.paragraphs[5].runs[10].text[0] == " ":
       myDoc.paragraphs[5].runs[10].text = myDoc.paragraphs[5].runs[10].text[1:]    
    newID = identifer[iPos][-6:]
    myDoc.paragraphs[2].runs[4].text = myDoc.paragraphs[2].runs[4].text+"esklass: "+klass
    myDoc.paragraphs[2].runs[0].text = myDoc.paragraphs[2].runs[0].text+"mer: "+newID 
    parLength = len(myDoc.paragraphs)
    rangeList=list(range(0,parLength))
    rangeList.reverse() 
    for indx in rangeList:
        if len(myDoc.paragraphs[indx].text) == 0:
            delete_paragraph(myDoc.paragraphs[indx])
        else:
            break
    print(f"{Title}.docx")
    Title = Title.replace(" ","")
    myDoc.save(f"C:/myGit/TestBox/SpltDoc/{Title}.docx")
    pfPos +=1
    iPos +=1
    
docList=os.listdir(r"C:\myGit\TestBox\SpltDoc")
os.chdir(r"C:\myGit\TestBox\SpltDoc")


for file in docList:
    myStr=f"soffice.exe --headless --convert-to pdf C:/myGit/TestBox/SpltDoc/{file}"
    p1=subprocess.run(myStr, shell=True)


#Create new "Big" document. 
from docx import Document

docList = list(filter(lambda x: ".docx" in x, os.listdir("C:/myGit/TestBox/SpltDoc")))

def combine_word_documents(files):
    merged_document = Document()
    for index, file in enumerate(files):
        sub_doc = Document(file)
        # Don't add a page break if you've reached the last file.
        if index < len(files)-1:
           sub_doc.add_page_break()
        for element in sub_doc.element.body:
            merged_document.element.body.append(element)
    merged_document.save('merged.docx')

combine_word_documents(docList)