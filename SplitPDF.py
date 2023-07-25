# -*- coding: utf-8 -*-
"""
Created on Mon Nov 18 21:19:58 2019
@author: MartinTropse
"""

import pytesseract 
from wand.image import Image as wi
from PIL import Image
import io
import re
import docx
import time
import os


pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

os.chdir(r"C:\myGit\TestBox")
myPdf="C:/myGit/TestBox/Objekt1_20.pdf"



machineList= []
extracted_text = []    
imgBlobs = []
pdf = wi(filename=myPdf, resolution = 300)
pdfImg = pdf.convert('jpeg')
pdf.close()

for img in pdfImg.sequence:
    page=wi(image=img) #This opens an actual image
    imgBlobs.append(page.make_blob('jpeg')) #This object contains byte data, make_blob cause the conversion    
#The bytes object is the converted back to image before being passed to pytesseract    

for imgBlob in imgBlobs:
    im = Image.open(io.BytesIO(imgBlob)) #"im" is an PIL, JpegImageFile which is assumes is the prefered object of tessearct
    text = pytesseract.image_to_string(im, lang='swe')
    extracted_text.append(text)    
pos = -1

pdfString=str(extracted_text)
pat = re.compile(r'Nr Namn Klass\\n\d{6}', re.M) # pattern: Upercase, then anything that is not in (.!?), then one of them

#https://www.dataquest.io/wp-content/uploads/2019/03/python-regular-expressions-cheat-sheet.pdf

pdfaList=pat.split(pdfString)
identifer = re.findall(pat, pdfString)
headTitle = re.compile(r"(\w+\s)+\d")








ntrBsk=re.compile(r"\\nNaturvärdesbeskrivning\\n.+(!?\\nÅtgärdsbehov\\n)")
ntrBed=re.compile(r"\\nNaturvärdesbedömning\\n.+(=?\\nNaturvärdesbeskrivning\\n)")
åtBhv=re.compile(r"\\nÅtgärdsbehov\\n.+(=?\',\s\')")

bdmTitle = re.compile(r"\\nNaturvärdesbedömning\\n\\")
bskTitle = re.compile(r"\\nNaturvärdesbeskrivning\\n\\")
åtgTitle = re.compile(r"\\n\Åtgärdsbehov\\n\\")

iPos = 0
pfPos = 1
pageCount = 0
page = ""

for x in pdfaList:
    myDoc = docx.Document()
    docStr=pdfaList[pfPos]
    docStr+page
    pageCheck=re.search("den\s\d\s.+\d{4}\sSida.+", docStr)
    if not pageCheck:
        page += docStr
        continue
    page = ""
    docStr=re.sub("den\s\d\s.+\d{4}\sSida.+", "", docStr)
    Title_0=re.search(headTitle, docStr)
    mainTitle=myDoc.add_paragraph().add_run(Title_0.group(0)[:-1])
    titleFont = mainTitle.font
    titleFont.bold = True
    titleFont = docx.shared.Pt(30)
    klass=Title_0.group(0)[-1]
    outBdm=re.search(ntrBed, docStr)
    Title_1=re.search(bdmTitle, docStr)
    bdmRun1=myDoc.add_paragraph().add_run(Title_1.group(0)[2:-3])
    bdmRun2=myDoc.add_paragraph().add_run(outBdm.group(0)[26:-26])
    font = bdmRun1.font
    font.bold = True
    font.name = 'Arial'
    font.size = docx.shared.Pt(16)   
    outBsk=re.search(ntrBsk, docStr)
    Title_2=re.search(bskTitle, docStr)        
    bskRun1=myDoc.add_paragraph().add_run(Title_2.group(0)[2:-3])
    font2 = bskRun1.font
    font2.bold = True
    font2.name = 'Arial'
    font2.size = docx.shared.Pt(16)    
    bskRun2=myDoc.add_paragraph().add_run(outBsk.group(0)[28:-18])
    outBhv=re.search(åtBhv, docStr)
    Title_3=re.search(åtgTitle, docStr)
    bhvRun1=myDoc.add_paragraph().add_run(Title_3.group(0)[2:-3])
    bhvRun2=myDoc.add_paragraph().add_run(outBhv.group(0)[18:])
    font3 = bhvRun1.font
    font3.bold = True
    font3.name = 'Arial'
    font3.size = docx.shared.Pt(16)
    newID = identifer[iPos][-6:]
    myDoc.save(f"{newID}.docx")
    pfPos +=1
    iPos +=1
    break
    

myDoc = docx.Document("3GreatDocument.docx")
imDoc.paragraphs[1].runs[0].text = "\n\t\tMyNewName!"

myDoc.paragraphs[5].runs[2].text #Naturvärdesbedömning text 
myDoc.paragraphs[5].runs[5].text #Naturvärdesbeskrivning text 
myDoc.paragraphs[5].runs[9].text #Åtgärdsbehov text 


for x in range(0, len(myDoc.paragraphs[5].runs)):
    print(x, imDoc.paragraphs[5].runs[x].text)

for x in range(0, len(myDoc.paragraphs[6].runs)):
    print(x, imDoc.paragraphs[6].runs[x].text)



myDoc.save("Test43.docx")

iPos = 0
pfPos = 1

ntrBsk=re.compile(r"\\nNaturvärdesbeskrivning\\n.+(!?\\nÅtgärdsbehov\\n)")
ntrBed=re.compile(r"\\nNaturvärdesbedömning\\n.+(=?\\nNaturvärdesbeskrivning\\n)")
åtBhv=re.compile(r"(\\nÅtgärdsbehov\\n.+(=?\',\s\')|Åtgärdsbehov\\n\\n.+")

åtBhv2=re.compile(r"Åtgärdsbehov\\n.+")

iPos = 0
pfPos = 1

#page = ""

iPos = 0
pfPos = 1

for x in pdfaList:
    myDoc = docx.Document("4GreatDocument.docx")
    docStr=pdfaList[pfPos]
    docStr=re.sub("den\s\d{1,2}\s[a-z]{1,14}\s\d{4}\sSida\s\d{1,3}\sav\s\d{1,3}", "", docStr)
    Title_0=re.search(headTitle, docStr)
    klass=Title_0.group(0)[-1]
    myDoc.paragraphs[1].runs[1].text = Title_0.group(0)[:-1]
    #myDoc.save(f"{newID}.docx")
    outBdm=re.search(ntrBed, docStr)
    myDoc.paragraphs[5].runs[2].text = outBdm.group(0)[26:-26].replace(r"\n", " ")+"\n"
    outBsk=re.search(ntrBsk, docStr)
    myDoc.paragraphs[5].runs[6].text = outBsk.group(0)[28:-18].replace(r"\n", " ")+"\n"
    outBhv=re.search(åtBhv, docStr)
    if outBhv:
        myDoc.paragraphs[5].runs[9].text = outBhv.group(0)[18:-5].replace(r"\n", " ")+"\n"
    else:
        outBhv=re.search(åtBhv2, docStr)
        print("RegEx2")
        myDoc.paragraphs[5].runs[9].text = outBhv.group(0)[14:-5].replace(r"\n", " ")+"\n"
    newID = identifer[iPos][-6:]
    myDoc.paragraphs[3].runs[0].text = newID+"\t\t"+klass
    print(f"{newID}.docx")
    myDoc.save(f"{newID}.docx")
    pfPos +=1
    iPos +=1



paragraph = myDoc.add_paragraph()
paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

paragraph=myDoc.paragraphs[0]




paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
paragraph.alignment
RIGHT (2)

#paragraph.alignment = None
#paragraph.alignment
#None


#document = Document(args.inputFile)

#changing the page margins
sections = myDoc.sections

for section in sections:
        section.left_margin = 2
#    section.top_margin = Cm(margin)
#    section.bottom_margin = Cm(margin)
#    section.right_margin = Cm(margin)

document.save(args.outputFile)

myDoc.save("a_doc.docx")

#outStr=re.sub(r"\\n", r"\\n\t\t", outBdm.group(0))
testStr=myDoc.paragraphs[0].text



newStr=re.sub(r"\\n", r"\\n\t\t",testStr)



os.chdir(r"C:\myGit\TestBox")
#Loop through all files and replace(\n). Yeah not great ^

docStr=re.sub(r"\n", "", docStr)

#imDoc.paragraphs[0].runs
#imDoc.paragraphs[3].text

imDoc=docx.Document("2GreatDocument.docx")

imDoc.save("3GreatDocument.docx")

for x in range(5, len(imDoc.paragraphs)):
    imDoc.paragraphs[x].remove()




paraObj2=imDoc.paragraphs[1].add_run("\n\t\tBrant SO Killingholmen")


imDoc.paragraphs[1].runs[0].text

font=paraObj2.font
font.bold = True
font.name = 'Arial'
font.size = docx.shared.Pt(32)

paraObj4=imDoc.paragraphs[4].add_run("\t\t192413 \t\t 3")
font=paraObj4.font
font.bold = True
font.name = 'Arial'
font.size = docx.shared.Pt(14)



imDoc.add_paragraph("I should have been an author!")
len(imDoc.paragraphs)






for x in range(5, len(imDoc.paragraphs)):
    delete_paragraph(imDoc.paragraphs[x])

myDoc.paragraphs[5].runs[7].text
myDoc.save("4GreatDocument.docx")

p = myDoc.paragraphs[5].runs[8]._element
p.getparent().remove(p)
p._p = p._element = None

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

























#
#path=baseFold+"/"+file[:-4]
#os.makedirs(path, exist_ok=True)
#for page in extracted_text:
#    pos += 1
#    newSent = ""  
#    splitSent=pat.findall(page)
#    for sentc in splitSent:
#        newSent += re.sub("\n", " ", sentc)
#    extracted_text[pos] = newSent
#
#
#nPage = 0
#
#for page in extracted_text:
#    chkVal = 0
#    nPage += 1
#    pat = re.compile(r'([A-Z,Å,Ä,Ö][^\.!?]*[\.!?])', re.M)
#    splitSent=pat.findall(page)
#    nLine = 0
#    for x in splitSent:
#        nLine +=1
#        mash=re.search('CO2|MWh|TWh|KWh',x, re.IGNORECASE)
#        if mash:
#            machineList.append(x+" Page number "+str(nPage)+" sentence "+str(nLine))
            
            #for x in pdfaList:
#    myDoc = docx.Document()
#    docStr=pdfaList[pfPos]
#    docStr+page
#    pageCheck=re.search("den\s\d\s.+\d{4}\sSida.+", docStr)
#    if not pageCheck:
#        page += docStr
#        continue
#    page = ""
